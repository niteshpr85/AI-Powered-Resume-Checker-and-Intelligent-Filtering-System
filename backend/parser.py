"""Resume parsing for PDF and DOCX files."""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Dict, List

import pdfplumber
from docx import Document
from PyPDF2 import PdfReader

from backend.skill_extractor import extract_skills
from utils.constants import EDUCATION_KEYWORDS, EDUCATION_RANK
from utils.text_cleaner import clean_text


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    plumber_pages: List[str] = []
    pypdf_pages: List[str] = []

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                plumber_pages.append(page.extract_text() or "")
    except Exception:
        plumber_pages = []

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            pypdf_pages.append(page.extract_text() or "")
    except Exception:
        pypdf_pages = []

    plumber_text = "\n".join(plumber_pages).strip()
    pypdf_text = "\n".join(pypdf_pages).strip()

    # Some resumes parse better with one extractor than the other.
    return plumber_text if len(plumber_text) >= len(pypdf_text) else pypdf_text


def _extract_text_from_docx(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    parts: List[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_name(text: str, fallback: str) -> str:
    for line in text.splitlines():
        clean_line = line.strip()
        if not clean_line:
            continue
        lower = clean_line.lower()
        if "@" in clean_line:
            continue
        if re.search(r"\+?\d[\d\-\s]{7,}", clean_line):
            continue
        if "linkedin" in lower or "github" in lower or "http" in lower:
            continue
        if len(clean_line.split()) <= 6:
            return clean_line
        break
    return fallback


def _extract_experience_years(text: str) -> int:
    patterns = [
        r"(\d{1,2}(?:\.\d)?)\+?\s+years?\s+of\s+experience",
        r"experience\s*[:\-]?\s*(\d{1,2}(?:\.\d)?)\+?\s+years?",
        r"(\d{1,2}(?:\.\d)?)\+?\s*yrs?",
    ]
    years: List[float] = []
    lower = text.lower()
    for pattern in patterns:
        years.extend(float(match) for match in re.findall(pattern, lower))
    if not years:
        return 0
    return min(int(max(years)), 40)


def _extract_education_level(text: str) -> str:
    lower = text.lower()
    detected = "high school"
    best_rank = 1
    for level, terms in EDUCATION_KEYWORDS.items():
        for term in terms:
            if term in lower and EDUCATION_RANK[level] > best_rank:
                detected = level
                best_rank = EDUCATION_RANK[level]
    return detected


def parse_resume_file(uploaded_file) -> Dict[str, object]:
    """Parse a Streamlit UploadedFile and return structured candidate data."""
    file_name = uploaded_file.name
    suffix = Path(file_name).suffix.lower()
    file_bytes = uploaded_file.getvalue()

    if suffix == ".pdf":
        raw_text = _extract_text_from_pdf(file_bytes)
    elif suffix == ".docx":
        raw_text = _extract_text_from_docx(file_bytes)
    else:
        raw_text = file_bytes.decode("utf-8", errors="ignore")

    cleaned = clean_text(raw_text)
    fallback_name = Path(file_name).stem.replace("_", " ").strip().title()
    candidate_name = _extract_name(cleaned, fallback_name)
    warning = ""
    if len(cleaned) < 180:
        warning = (
            "Low extracted text. This file may be image-based/scanned. "
            "Use a text-based PDF or DOCX for accurate parsing."
        )

    return {
        "name": candidate_name,
        "file_name": file_name,
        "raw_text": cleaned,
        "text_length": len(cleaned),
        "parse_warning": warning,
        "skills": extract_skills(cleaned),
        "experience_years": _extract_experience_years(cleaned),
        "education": _extract_education_level(cleaned),
    }
